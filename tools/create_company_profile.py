from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PUBLIC = ROOT / "public"
OUTPUT_DIR = ROOT / "deliverables"
OUTPUT = OUTPUT_DIR / "SpaceChamps_Company_Profile_2026.docx"

# narrative_proposal preset, resolved to exact values. Brand-colored components are
# named overrides used consistently throughout the profile.
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120

NAVY = "071B33"
NAVY_2 = "0B294A"
BLUE = "176BFF"
CYAN = "21A8F6"
ICE = "EAF6FF"
PALE = "F4F8FC"
MIST = "E5EEF7"
SLATE = "52677D"
INK = "14263A"
WHITE = "FFFFFF"
LINE = "CAD8E6"
GREEN = "1A8F70"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=LINE, size=6, inside=True) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    sides = ["top", "left", "bottom", "right"]
    if inside:
        sides += ["insideH", "insideV"]
    for side in sides:
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths_dxa):
            cell = row.cells[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(
                cell,
                CELL_TOP_BOTTOM_DXA,
                CELL_START_END_DXA,
                CELL_TOP_BOTTOM_DXA,
                CELL_START_END_DXA,
            )


def set_font(run, size=None, color=INK, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border(paragraph, color=CYAN, size=12, side="left", space=8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_hyperlink(paragraph, text: str, url: str, color=BLUE, underline=True):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE ")
    set_font(run, 8.5, SLATE, bold=True)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_kicker(doc, text: str, after=5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper())
    set_font(run, 8.5, CYAN, bold=True)
    run.font.letter_spacing = Pt(1.1) if hasattr(run.font, "letter_spacing") else None
    return paragraph


def add_section_title(doc, title: str, subtitle: str | None = None):
    add_kicker(doc, "SpaceChamps company profile")
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(title)
    set_font(run, 23, NAVY, bold=True)
    if subtitle:
        lead = doc.add_paragraph()
        lead.paragraph_format.space_before = Pt(0)
        lead.paragraph_format.space_after = Pt(12)
        lead.paragraph_format.line_spacing = 1.2
        lead.paragraph_format.keep_with_next = True
        r = lead.add_run(subtitle)
        set_font(r, 11.5, SLATE)
    return paragraph


def add_body(doc, text: str, after=8, bold_lead: str | None = None, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.333
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r1 = paragraph.add_run(bold_lead)
        set_font(r1, 11, INK, bold=True)
        r2 = paragraph.add_run(text[len(bold_lead):])
        set_font(r2, 11, INK)
    else:
        run = paragraph.add_run(text)
        set_font(run, 11, INK)
    return paragraph


def add_callout(doc, label: str, text: str, fill=PALE, accent=CYAN):
    table = doc.add_table(rows=1, cols=1)
    # Named callout override: the 240 DXA indent matches the callout's 240 DXA
    # start cell margin so its visible content aligns with surrounding text.
    set_table_geometry(table, [CONTENT_WIDTH_DXA], indent_dxa=240)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=180, start=240, bottom=180, end=240)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    set_paragraph_border(paragraph, accent, 18, "left", 12)
    r1 = paragraph.add_run(label.upper() + "  ")
    set_font(r1, 9, accent, bold=True)
    r2 = paragraph.add_run(text)
    set_font(r2, 11, INK, bold=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(0)
    return table


def add_image(doc, path: Path, width: float, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, alt=""):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    if alt:
        doc_pr = shape._inline.docPr
        doc_pr.set("descr", alt)
        doc_pr.set("title", alt)
    return shape


def add_table_heading(cell, text: str, color=WHITE, size=9.5):
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text.upper())
    set_font(run, size, color, bold=True)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_cell_text(cell, title: str, description: str | None = None, title_color=INK, size=9.7):
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1 if description else 0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(title)
    set_font(run, size, title_color, bold=True)
    if description:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.1
        r2 = p2.add_run(description)
        set_font(r2, size - 0.4, SLATE)


def add_picture_to_cell(cell, path: Path, width: float, alt: str):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_document_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = rgb(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = rgb(BLUE)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = rgb(NAVY_2)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def add_running_furniture(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.space_before = Pt(0)
    set_paragraph_border(paragraph, LINE, 4, "bottom", 4)
    left = paragraph.add_run("SPACECHAMPS")
    set_font(left, 8.5, NAVY, bold=True)
    right = paragraph.add_run("  |  COMPANY PROFILE 2026")
    set_font(right, 8.5, SLATE, bold=False)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [7000, 2360], indent_dxa=0)
    remove_table_borders(table)
    p_left = table.cell(0, 0).paragraphs[0]
    p_left.paragraph_format.space_after = Pt(0)
    r = p_left.add_run("Precision in every point.")
    set_font(r, 8.5, SLATE, italic=True)
    p_right = table.cell(0, 1).paragraphs[0]
    add_page_number(p_right)

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("SPACECHAMPS  |  COMPANY PROFILE 2026")
    set_font(r, 8.5, SLATE, bold=True)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def add_cover(doc: Document) -> None:
    add_image(
        doc,
        PUBLIC / "spacechamps-brand-header.png",
        width=2.5,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        after=2,
        alt="SpaceChamps logo",
    )
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(6)
    r = kicker.add_run("COMPANY PROFILE  |  2026")
    set_font(r, 9, CYAN, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.line_spacing = 0.96
    r1 = title.add_run("Advanced Drone, LiDAR &\n")
    set_font(r1, 29, NAVY, bold=True)
    r2 = title.add_run("Geospatial Intelligence")
    set_font(r2, 29, BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.line_spacing = 1.15
    r = subtitle.add_run("Precision data for smarter decisions and stronger infrastructure.")
    set_font(r, 12.5, SLATE)

    add_image(
        doc,
        PUBLIC / "spacechamps-hero.png",
        width=6.5,
        after=12,
        alt="Drone surveying mountainous infrastructure with LiDAR",
    )

    strap = doc.add_paragraph()
    strap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    strap.paragraph_format.space_before = Pt(0)
    strap.paragraph_format.space_after = Pt(4)
    r = strap.add_run("CAPTURE REALITY. DELIVER CERTAINTY.")
    set_font(r, 10.5, NAVY, bold=True)
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.paragraph_format.space_before = Pt(0)
    foot.paragraph_format.space_after = Pt(0)
    r = foot.add_run("Jakarta, Indonesia  |  Doha, Qatar  |  Southeast Asia & Middle East")
    set_font(r, 9.5, SLATE)


def add_company_overview(doc: Document) -> None:
    add_section_title(
        doc,
        "Precision data. Smarter decisions.",
        "SpaceChamps turns complex terrain and infrastructure into decision-ready geospatial intelligence.",
    )
    add_callout(
        doc,
        "Our purpose",
        "Make high-precision geospatial data accessible, reliable, and actionable for the projects that shape our world.",
    )
    add_body(
        doc,
        "Founded in 2012, SpaceChamps is a drone-based LiDAR, aerial mapping, and geospatial intelligence company serving clients across Southeast Asia and the Middle East. We combine field expertise, advanced sensors, RTK positioning, and AI-powered processing to convert real-world conditions into clear, usable project intelligence.",
    )
    add_body(
        doc,
        "Our work supports planning, design, construction, inspection, monitoring, and asset management. From a single site survey to a long-range regional mapping mission, every engagement is designed around accuracy, operational safety, fast turnaround, and practical deliverables.",
    )

    table = doc.add_table(rows=2, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    set_table_borders(table, MIST, 7)
    numbers = ["14+", "500+", "250+", "98.7%"]
    labels = ["Years of experience", "Projects delivered", "Enterprise clients", "Data quality target"]
    for i, value in enumerate(numbers):
        set_cell_shading(table.cell(0, i), NAVY if i % 2 == 0 else NAVY_2)
        p = table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, 22, WHITE, bold=True)
        table.cell(0, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(table.cell(1, i), PALE)
        p2 = table.cell(1, i).paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(labels[i])
        set_font(r2, 9, SLATE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("MISSION")
    set_font(r, 9, CYAN, bold=True)
    add_body(
        doc,
        "To deliver accurate, reliable, and actionable geospatial data that helps organizations plan smarter, build faster, and operate with confidence.",
        after=5,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("VISION")
    set_font(r, 9, CYAN, bold=True)
    add_body(
        doc,
        "To set the standard for field-to-insight geospatial intelligence across fast-growing markets.",
        after=0,
    )


SERVICES = [
    ("LiDAR Data Acquisition", "High-density point clouds and precise elevation data for complex sites."),
    ("Topographical Surveys", "Terrain, contour, and engineering survey deliverables for design teams."),
    ("Aerial Photogrammetry", "High-resolution imagery, orthomosaics, and 3D reconstruction."),
    ("DTM & DSM Models", "Detailed terrain and surface models for analysis, engineering, and planning."),
    ("Orthomosaic Mapping", "Georeferenced, measurable maps for clear site intelligence."),
    ("Corridor Surveys", "Linear mapping for roads, railways, pipelines, and powerline routes."),
    ("Construction Monitoring", "Repeatable progress tracking, change detection, and volumetrics."),
    ("Mining & Quarry Surveys", "Stockpile volumes, pit analysis, haul roads, and operational monitoring."),
    ("Agriculture & Environment", "Crop assessment, land analysis, and environmental mapping."),
    ("Infrastructure Inspection", "Safe aerial inspection of structures, corridors, and critical assets."),
    ("Long-Endurance VTOL", "Extended-range missions for remote locations and large-area projects."),
    ("Deployment & Support", "Mission planning, field operations, QA, delivery, and technical support."),
]


def add_services(doc: Document) -> None:
    add_section_title(
        doc,
        "Comprehensive geospatial solutions",
        "A complete service portfolio spanning field capture, mapping, modelling, monitoring, inspection, and delivery.",
    )
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3000, 6360])
    set_table_borders(table, LINE, 6)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, NAVY)
    add_table_heading(table.cell(0, 0), "Capability")
    add_table_heading(table.cell(0, 1), "What it delivers")
    set_repeat_table_header(table.rows[0])
    for index, (title, description) in enumerate(SERVICES):
        cells = table.add_row().cells
        set_table_geometry(table, [3000, 6360])
        fill = WHITE if index % 2 == 0 else PALE
        for cell in cells:
            set_cell_shading(cell, fill)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_cell_text(cells[0], title, title_color=NAVY, size=9.6)
        add_cell_text(cells[1], description, title_color=SLATE, size=9.4)
    add_callout(
        doc,
        "Flexible engagement",
        "Services can be delivered independently or combined into a complete capture-to-delivery program.",
        fill=ICE,
        accent=BLUE,
    )


INDUSTRIES = [
    ("Government & public agencies", "Base mapping, land administration, planning, disaster readiness, and public asset intelligence."),
    ("Engineering & design firms", "Survey-grade inputs for feasibility, concept design, detailed design, and verification."),
    ("Construction companies", "Progress monitoring, cut-and-fill analysis, as-built documentation, and site coordination."),
    ("Mining & quarry operators", "Pit mapping, stockpile volumes, haul-road monitoring, and operational decision support."),
    ("Agriculture & environment", "Crop and land analysis, environmental mapping, change detection, and site monitoring."),
    ("Infrastructure developers", "Corridor mapping, inspection, asset models, and lifecycle intelligence for critical networks."),
]


def add_industries(doc: Document) -> None:
    add_section_title(
        doc,
        "Built for mission-critical sectors",
        "Our workflows are designed for organizations that depend on accurate site intelligence to make high-consequence decisions.",
    )
    add_image(
        doc,
        PUBLIC / "equipment" / "matrice-350-field.jpg",
        width=4.9,
        after=10,
        alt="DJI Matrice 350 RTK conducting an aerial survey",
    )
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3000, 6360])
    set_table_borders(table, LINE, 6)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, NAVY)
    add_table_heading(table.cell(0, 0), "Sector")
    add_table_heading(table.cell(0, 1), "Typical applications")
    set_repeat_table_header(table.rows[0])
    for index, (sector, applications) in enumerate(INDUSTRIES):
        cells = table.add_row().cells
        set_table_geometry(table, [3000, 6360])
        fill = WHITE if index % 2 == 0 else PALE
        for cell in cells:
            set_cell_shading(cell, fill)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_cell_text(cells[0], sector, title_color=NAVY, size=9.6)
        add_cell_text(cells[1], applications, title_color=SLATE, size=9.35)
    add_body(
        doc,
        "Across every sector, the objective is the same: reduce uncertainty, strengthen coordination, and give decision-makers a reliable digital view of the physical world.",
        after=0,
    )


EQUIPMENT = [
    (
        "matrice-350-field.jpg",
        "DJI Matrice 350 RTK",
        "Enterprise UAV",
        "All-weather field operations | RTK positioning | up to 55-minute flight time | dual-battery workflow",
    ),
    (
        "zenmuse-l2-field.jpg",
        "DJI Zenmuse L2",
        "LiDAR payload",
        "Multi-return scanning | integrated RGB imaging | precise point-cloud capture | terrain mapping",
    ),
    (
        "long-endurance-vtol-field.jpg",
        "Long-Endurance VTOL",
        "Large-area mapping",
        "4+ hour endurance | 500+ sq km mission potential | runway-free operation | autonomous launch and recovery",
    ),
    (
        "trimble-r12i-field.jpg",
        "Trimble R12i",
        "GNSS RTK receiver",
        "IMU tilt compensation | multi-constellation tracking | centimeter-level positioning | all-weather use",
    ),
    (
        "trimble-tdc6-field.jpg",
        "Trimble TDC6",
        "Mobile GIS & field data",
        "Rugged field collection | connected workflows | sunlight-readable display | mobile GIS delivery",
    ),
]


def add_technology(doc: Document) -> None:
    add_section_title(
        doc,
        "Advanced equipment. Intelligent workflows.",
        "A field-proven technology stack selected for accuracy, reliability, and operational flexibility.",
    )
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2100, 2500, 4760])
    set_table_borders(table, LINE, 6)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, NAVY)
    add_table_heading(table.cell(0, 0), "Platform")
    add_table_heading(table.cell(0, 1), "Role")
    add_table_heading(table.cell(0, 2), "Operational value")
    set_repeat_table_header(table.rows[0])
    for index, (filename, name, role, value) in enumerate(EQUIPMENT):
        cells = table.add_row().cells
        set_table_geometry(table, [2100, 2500, 4760])
        fill = WHITE if index % 2 == 0 else PALE
        for cell in cells:
            set_cell_shading(cell, fill)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_picture_to_cell(
            cells[0],
            PUBLIC / "equipment" / filename,
            width=1.23,
            alt=name,
        )
        add_cell_text(cells[1], name, role, title_color=NAVY, size=9.5)
        add_cell_text(cells[2], value, title_color=SLATE, size=9.1)
    add_callout(
        doc,
        "Technology principle",
        "Use the right platform, sensor, positioning method, and processing workflow for the mission - not a one-size-fits-all package.",
        fill=ICE,
        accent=CYAN,
    )


WORKFLOW = [
    ("01", "Mission planning", "Define objectives, accuracy, site constraints, flight paths, control strategy, and safety protocols."),
    ("02", "Drone deployment", "Capture LiDAR, high-resolution imagery, thermal data, or GNSS control using the selected platform."),
    ("03", "Field quality control", "Validate coverage, positioning, sensor health, and data completeness before leaving site."),
    ("04", "Cloud processing", "Generate point clouds, imagery products, terrain models, and classified datasets."),
    ("05", "AI analysis & QA", "Classify features, extract measurements, flag anomalies, and verify outputs against specifications."),
    ("06", "Delivery & support", "Package decision-ready maps, models, GIS files, and reports with technical handover support."),
]


def add_workflow(doc: Document) -> None:
    add_section_title(
        doc,
        "From takeoff to decision-ready delivery",
        "A controlled, end-to-end workflow keeps quality, speed, and accountability visible at every stage.",
    )
    add_image(
        doc,
        PUBLIC / "equipment" / "long-endurance-vtol-field.jpg",
        width=3.9,
        after=9,
        alt="Long-endurance VTOL platform prepared for a large-area survey",
    )
    table = doc.add_table(rows=0, cols=3)
    for index, (number, title, description) in enumerate(WORKFLOW):
        cells = table.add_row().cells
        set_table_geometry(table, [900, 2500, 5960])
        fill = PALE if index % 2 == 0 else WHITE
        for cell in cells:
            set_cell_shading(cell, fill)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(number)
        set_font(r0, 13, CYAN, bold=True)
        add_cell_text(cells[1], title, title_color=NAVY, size=9.6)
        add_cell_text(cells[2], description, title_color=SLATE, size=9.2)
    set_table_borders(table, LINE, 5, inside=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("TYPICAL DELIVERABLES")
    set_font(r, 9, CYAN, bold=True)
    add_body(
        doc,
        "Point clouds | orthomosaics | DTM and DSM | contours and CAD-ready survey data | 3D models | GIS layers | volumetric reports | inspection records | web-ready visualizations",
        after=0,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )


QUALITY = [
    ("Survey control", "RTK/GNSS positioning, ground control, and mission-specific accuracy planning."),
    ("Field assurance", "On-site checks for coverage, sensor health, image quality, and completeness."),
    ("Processing assurance", "Automated classification, feature extraction, and repeatable QA checkpoints."),
    ("Delivery assurance", "Specification review, documented outputs, and technical handover support."),
]


def add_quality(doc: Document) -> None:
    add_section_title(
        doc,
        "Quality engineered into every stage",
        "Accuracy is not a final check. It is designed into planning, capture, processing, analysis, and delivery.",
    )

    image_table = doc.add_table(rows=1, cols=2)
    # Named figure-pair override: the 70 DXA indent matches the figure cell inset.
    set_table_geometry(image_table, [4680, 4680], indent_dxa=70)
    remove_table_borders(image_table)
    for cell in image_table.rows[0].cells:
        set_cell_margins(cell, top=0, start=70, bottom=80, end=70)
    add_picture_to_cell(
        image_table.cell(0, 0),
        PUBLIC / "equipment" / "trimble-r12i-field.jpg",
        width=2.55,
        alt="Surveyor using a Trimble R12i GNSS receiver",
    )
    add_picture_to_cell(
        image_table.cell(0, 1),
        PUBLIC / "equipment" / "trimble-tdc6-field.jpg",
        width=2.55,
        alt="Trimble TDC6 mobile GIS field data collector",
    )
    for index, caption in enumerate(("GNSS control & verification", "Connected field data workflows")):
        p = image_table.cell(0, index).add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(caption)
        set_font(r, 8.8, SLATE, bold=True)

    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2800, 6560])
    set_table_borders(table, LINE, 6)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, NAVY)
    add_table_heading(table.cell(0, 0), "Control layer")
    add_table_heading(table.cell(0, 1), "How quality is protected")
    set_repeat_table_header(table.rows[0])
    for index, (control, method) in enumerate(QUALITY):
        cells = table.add_row().cells
        set_table_geometry(table, [2800, 6560])
        fill = WHITE if index % 2 == 0 else PALE
        for cell in cells:
            set_cell_shading(cell, fill)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_cell_text(cells[0], control, title_color=NAVY, size=9.5)
        add_cell_text(cells[1], method, title_color=SLATE, size=9.25)

    add_callout(
        doc,
        "AI with human oversight",
        "Automation accelerates classification and feature extraction, while experienced specialists remain accountable for final quality.",
        fill=ICE,
        accent=GREEN,
    )


VALUES = [
    ("Precision", "Centimeter-level thinking in every dataset, decision, and deliverable."),
    ("Innovation", "Practical adoption of advanced hardware, AI, and cloud workflows."),
    ("Reliability", "Consistent quality, clear communication, and disciplined delivery."),
    ("Partnership", "A collaborative approach that works as an extension of the client team."),
]

JOURNEY = [
    ("2012", "SpaceChamps founded by survey professionals and drone pilots."),
    ("2015", "First international expansion into multi-region project delivery."),
    ("2018", "Advanced LiDAR sensors and RTK positioning integrated into operations."),
    ("2021", "Cloud-based AI processing launched for classification and feature extraction."),
    ("2024", "500+ delivered projects milestone reached."),
    ("2026", "14 years of operations serving clients across Southeast Asia and the Middle East."),
]


def add_values_and_journey(doc: Document) -> None:
    add_section_title(
        doc,
        "Built on precision and partnership",
        "The company has grown by combining technical discipline with a straightforward commitment to client outcomes.",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("CORE VALUES")
    set_font(r, 9, CYAN, bold=True)
    values_table = doc.add_table(rows=1, cols=2)
    set_table_geometry(values_table, [2500, 6860])
    set_table_borders(values_table, LINE, 6)
    for cell in values_table.rows[0].cells:
        set_cell_shading(cell, NAVY)
    add_table_heading(values_table.cell(0, 0), "Value")
    add_table_heading(values_table.cell(0, 1), "What it means in practice")
    set_repeat_table_header(values_table.rows[0])
    for index, (value, meaning) in enumerate(VALUES):
        cells = values_table.add_row().cells
        set_table_geometry(values_table, [2500, 6860])
        fill = WHITE if index % 2 == 0 else PALE
        for cell in cells:
            set_cell_shading(cell, fill)
        add_cell_text(cells[0], value, title_color=NAVY, size=9.6)
        add_cell_text(cells[1], meaning, title_color=SLATE, size=9.3)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("OUR JOURNEY")
    set_font(r, 9, CYAN, bold=True)
    journey_table = doc.add_table(rows=0, cols=2)
    for index, (year, milestone) in enumerate(JOURNEY):
        cells = journey_table.add_row().cells
        set_table_geometry(journey_table, [1400, 7960])
        fill = ICE if index % 2 == 0 else WHITE
        for cell in cells:
            set_cell_shading(cell, fill)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p_year = cells[0].paragraphs[0]
        p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_year.paragraph_format.space_after = Pt(0)
        r_year = p_year.add_run(year)
        set_font(r_year, 12, BLUE, bold=True)
        add_cell_text(cells[1], milestone, title_color=INK, size=9.25)
    set_table_borders(journey_table, LINE, 5)

    add_callout(
        doc,
        "Regional footprint",
        "Operational presence in Jakarta and Doha, with a service focus across Southeast Asia and the Middle East.",
        fill=PALE,
        accent=BLUE,
    )


def add_contact(doc: Document) -> None:
    add_image(
        doc,
        PUBLIC / "spacechamps-hero.png",
        width=6.5,
        after=13,
        alt="Drone LiDAR mapping over regional infrastructure",
    )
    add_kicker(doc, "Let's work together")
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.line_spacing = 1.0
    r1 = title.add_run("Ready to map what matters next?")
    set_font(r1, 25, NAVY, bold=True)
    add_body(
        doc,
        "Tell us about your site, project objectives, required accuracy, timeline, and preferred deliverables. The SpaceChamps team will help define the right field and processing approach.",
        after=10,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [2300, 7060])
    set_table_borders(table, LINE, 6)
    contact_rows = [
        ("WhatsApp", "+62 838-5209-4053", "https://wa.me/6283852094053"),
        ("Email", "engwarsame16@gmail.com", "mailto:engwarsame16@gmail.com"),
        ("Website", "spacechamps.vercel.app", "https://spacechamps.vercel.app"),
        ("Locations", "Jakarta, Indonesia  |  Doha, Qatar", None),
    ]
    for index, (label, value, url) in enumerate(contact_rows):
        label_cell, value_cell = table.rows[index].cells
        set_cell_shading(label_cell, NAVY if index % 2 == 0 else NAVY_2)
        set_cell_shading(value_cell, PALE if index % 2 == 0 else WHITE)
        add_cell_text(label_cell, label.upper(), title_color=WHITE, size=9.5)
        paragraph = value_cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        if url:
            add_hyperlink(paragraph, value, url, color=BLUE, underline=False)
        else:
            run = paragraph.add_run(value)
            set_font(run, 10, INK, bold=True)
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(15)
    closing.paragraph_format.space_after = Pt(4)
    r = closing.add_run("PRECISION IN EVERY POINT.")
    set_font(r, 11, BLUE, bold=True)
    rights = doc.add_paragraph()
    rights.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rights.paragraph_format.space_after = Pt(0)
    r = rights.add_run("© 2026 SpaceChamps. All rights reserved.")
    set_font(r, 8.5, SLATE)


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "SpaceChamps Company Profile 2026"
    props.subject = "Advanced drone, LiDAR, and geospatial intelligence company profile"
    props.author = "SpaceChamps"
    props.keywords = "SpaceChamps, LiDAR, drone surveying, geospatial intelligence, aerial mapping, GIS"
    props.comments = "Prepared from the approved SpaceChamps website and brand materials."


def build() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_page(doc.sections[0])
    set_document_styles(doc)
    add_running_furniture(doc.sections[0])
    set_core_properties(doc)

    add_cover(doc)
    page_break(doc)
    add_company_overview(doc)
    page_break(doc)
    add_services(doc)
    page_break(doc)
    add_industries(doc)
    page_break(doc)
    add_technology(doc)
    page_break(doc)
    add_workflow(doc)
    page_break(doc)
    add_quality(doc)
    page_break(doc)
    add_values_and_journey(doc)
    page_break(doc)
    add_contact(doc)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
