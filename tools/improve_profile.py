"""
Improve SpaceChamps Company Profile document.
- Delete website row from contact table
- Polish and enrich content throughout
"""
from docx import Document
from copy import deepcopy
from lxml import etree

SRC = r"D:\spacechamps\deliverables\SpaceChamps_Company_Profile_2026.docx"
OUT = r"D:\spacechamps\deliverables\SpaceChamps_Company_Profile_2026.docx"

doc = Document(SRC)

# ============================================================
# 1. DELETE WEBSITE ROW FROM CONTACT TABLE (Table 14)
# ============================================================
contact_tbl = doc.tables[14]
tbl_elem = contact_tbl._tbl
# Find the row with WEBSITE text
rows_to_remove = []
for ri, row in enumerate(contact_tbl.rows):
    if "WEBSITE" in row.cells[0].text.upper():
        rows_to_remove.append(row._tr)
for tr in rows_to_remove:
    tbl_elem.remove(tr)
print("Deleted WEBSITE row from contact table.")

# ============================================================
# 2. IMPROVE PARAGRAPH CONTENT
# ============================================================
para_improvements = {
    # Tagline
    3: "Precision data for smarter decisions, stronger infrastructure, and faster project delivery.",
    # Mission statement - more powerful
    15: "To deliver accurate, reliable, and actionable geospatial intelligence that helps organizations plan smarter, build faster, and operate with total confidence in their data.",
    # Vision - more ambitious
    17: "To become the most trusted name in field-to-insight geospatial intelligence across emerging markets and fast-growing economies.",
    # Company description paragraph 1 - richer
    12: "Founded in 2012, SpaceChamps is a drone-based LiDAR, aerial mapping, and geospatial intelligence company serving clients across Southeast Asia and the Middle East. We combine deep field expertise, industry-leading sensors, RTK positioning, and AI-powered processing to transform real-world conditions into clear, decision-ready project intelligence.",
    # Company description paragraph 2 - stronger
    13: "Our work supports planning, design, construction, inspection, monitoring, and asset management at every stage. From a single-site survey to a long-range regional mapping mission, every engagement is built around accuracy, operational safety, rapid turnaround, and practical deliverables that teams can use immediately.",
    # Services intro - more benefit-driven
    21: "A complete service portfolio spanning field capture, mapping, modelling, monitoring, inspection, and delivery — each capability available independently or combined into a single end-to-end program.",
    # Sectors intro - more compelling
    26: "Our workflows are purpose-built for organizations that depend on accurate, timely site intelligence to make high-stakes decisions under real-world pressure.",
    # Sectors summary - stronger close
    28: "Across every sector, the objective remains the same: reduce uncertainty, strengthen coordination, and give decision-makers a reliable digital view of the physical world — so projects move forward with confidence.",
    # Technology intro
    32: "A field-proven technology stack carefully selected for centimeter-level accuracy, hardware reliability, and operational flexibility across diverse environments and mission profiles.",
    # Pipeline intro
    37: "A controlled, end-to-end workflow that keeps quality, speed, and accountability visible at every stage — from initial planning through to final handover.",
    # Quality intro
    44: "Accuracy is not a final checkpoint. It is engineered into every phase — planning, capture, processing, analysis, and delivery — with repeatable QA controls and documented accountability throughout.",
    # Values intro
    49: "SpaceChamps has grown by combining technical discipline with an uncompromising commitment to client outcomes — delivering work that teams can trust, on time, every time.",
    # CTA
    57: "Tell us about your site, project objectives, required accuracy, timeline, and preferred deliverables. The SpaceChamps team will respond within one business day with a proposed scope, crew plan, and processing approach tailored to your mission.",
    # Deliverables line
    40: "Point clouds | orthomosaics | DTM and DSM | contours and CAD-ready survey data | 3D reality models | GIS layers | volumetric reports | inspection records | interactive web visualizations",
    # Footer
    59: "\u00a9 2026 SpaceChamps. All rights reserved.",
}

for idx, new_text in para_improvements.items():
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        if p.runs:
            # Preserve first run formatting, replace text
            p.runs[0].text = new_text
            # Clear extra runs
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.text = new_text
        print(f"Updated paragraph {idx}: {new_text[:60]}...")

# ============================================================
# 3. ENRICH TABLE CELL CONTENT
# ============================================================
table_cell_improvements = {
    # Table 2 - Services (enrich descriptions)
    (2, 1, 1): "High-density point clouds and precise elevation data for complex, vegetated, or hard-to-access sites.",
    (2, 2, 1): "Terrain, contour, and engineering survey deliverables that give design teams confidence from day one.",
    (2, 3, 1): "High-resolution imagery, orthomosaics, and photogrammetric 3D reconstruction with measurable detail.",
    (2, 4, 1): "Detailed terrain and surface models engineered for analysis, design, planning, and earthwork decisions.",
    (2, 5, 1): "Georeferenced, measurable maps that deliver instant, clear site intelligence for any stakeholder.",
    (2, 6, 1): "Linear mapping for roads, railways, pipelines, and powerline routes — accurate over long distances.",
    (2, 7, 1): "Repeatable progress tracking, change detection, and volumetric reporting across project timelines.",
    (2, 8, 1): "Stockpile volumes, pit analysis, haul-road monitoring, and operational intelligence for active sites.",
    (2, 9, 1): "Crop health, land classification, environmental mapping, and change detection for sustainable decisions.",
    (2, 10, 1): "Safe, close-range aerial inspection of structures, corridors, and critical assets — no scaffolding required.",
    (2, 11, 1): "Extended-range missions for remote locations, large-area coverage, and time-critical regional projects.",
    (2, 12, 1): "Mission planning, field operations, QA, delivery, and ongoing technical support from a dedicated team.",

    # Table 4 - Sectors (enrich applications)
    (4, 1, 1): "Base mapping, land administration, urban planning, disaster readiness, and public asset intelligence at scale.",
    (4, 2, 1): "Survey-grade inputs for feasibility studies, concept design, detailed engineering, and construction verification.",
    (4, 3, 1): "Progress monitoring, cut-and-fill analysis, as-built documentation, and real-time site coordination.",
    (4, 4, 1): "Pit mapping, stockpile volumes, haul-road monitoring, and operational decision support for active extraction.",
    (4, 5, 1): "Crop and land analysis, environmental impact mapping, change detection, and ongoing site monitoring.",
    (4, 6, 1): "Corridor mapping, structural inspection, digital asset models, and lifecycle intelligence for critical networks.",

    # Table 9 - Quality controls (strengthen)
    (9, 1, 1): "RTK/GNSS positioning, ground control points, and mission-specific accuracy planning before takeoff.",
    (9, 2, 1): "On-site validation of coverage, sensor health, image quality, and data completeness before leaving the site.",
    (9, 3, 1): "Automated classification, feature extraction, and repeatable QA checkpoints with documented traceability.",
    (9, 4, 1): "Specification review, documented outputs, structured deliverables, and hands-on technical handover support.",

    # Table 11 - Values (make more distinctive)
    (11, 1, 1): "Centimeter-level thinking applied to every dataset, every decision, and every deliverable we produce.",
    (11, 2, 1): "Practical, early adoption of advanced hardware, AI, and cloud workflows that give clients a real edge.",
    (11, 3, 1): "Consistent quality, transparent communication, and disciplined delivery — project after project, year after year.",
    (11, 4, 1): "A collaborative approach that works as a seamless extension of the client's own team and objectives.",

    # Table 12 - Timeline (enrich milestones)
    (12, 0, 1): "SpaceChamps founded by survey professionals and drone pilots committed to a higher data standard.",
    (12, 1, 1): "First international expansion, establishing multi-region project delivery capabilities.",
    (12, 2, 1): "Advanced LiDAR sensors and RTK positioning integrated, raising the bar for field accuracy.",
    (12, 3, 1): "Cloud-based AI processing launched for automated classification and rapid feature extraction.",
    (12, 4, 1): "500+ delivered projects milestone reached across government, engineering, and enterprise sectors.",
    (12, 5, 1): "14 years of operations, serving clients across Southeast Asia and the Middle East with a growing fleet and team.",

    # Table 13 - Regional footprint callout
    (13, 0, 0): "REGIONAL FOOTPRINT  Operational headquarters in Jakarta, Indonesia and a branch office in Doha, Qatar — with active service coverage across Southeast Asia and the Middle East.",

    # Table 0 - Purpose callout
    (0, 0, 0): "OUR PURPOSE  Make high-precision geospatial data accessible, reliable, and actionable for the projects that shape our world.",

    # Table 3 - Flexible engagement callout
    (3, 0, 0): "FLEXIBLE ENGAGEMENT  Every service can be delivered independently or combined into a complete capture-to-delivery program tailored to the project scope.",

    # Table 6 - Technology principle callout
    (6, 0, 0): "TECHNOLOGY PRINCIPLE  Match the right platform, sensor, positioning method, and processing workflow to each mission — never a one-size-fits-all package.",

    # Table 10 - AI oversight callout
    (10, 0, 0): "AI WITH HUMAN OVERSIGHT  Automation accelerates classification and feature extraction, while experienced specialists remain fully accountable for final quality and accuracy.",
}

for (ti, ri, ci), new_text in table_cell_improvements.items():
    if ti < len(doc.tables):
        tbl = doc.tables[ti]
        if ri < len(tbl.rows) and ci < len(tbl.rows[ri].cells):
            cell = tbl.rows[ri].cells[ci]
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].text = new_text
                for run in cell.paragraphs[0].runs[1:]:
                    run.text = ""
            else:
                cell.paragraphs[0].text = new_text
            print(f"Updated T{ti} R{ri} C{ci}: {new_text[:60]}...")

# ============================================================
# SAVE
# ============================================================
doc.save(OUT)
print(f"\nSaved to: {OUT}")
